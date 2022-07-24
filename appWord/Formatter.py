import os

import socket
from docx import Document
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

import tasklist
from tasklist import *


class Formatter:
    settings: dict

    month = {"01": "февраля", "02": "января", "03": "марта", "04": "апреля", "05": "мая",
             "06": "июня", "07": "июля", "08": "августа", "09": "сентября", "10": "октября",
             "11": "ноября", "12": "декабря"}

    def __init__(self, path, settings, path_to_save):
        self.doc = Document(path)
        self.path = path
        self.name = ""
        self.number = ""
        self.settings = settings
        self.path_to_save = path_to_save
        self.path_to_save_dif = f"{self.path_to_save}\\Разное"
        self.Dates = []
        self.Times = []

    def create_element(self, name):
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(self, run):
        fldStart = self.create_element('w:fldChar')
        self.create_attribute(fldStart, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'separate')

        fldChar2 = self.create_element('w:t')
        fldChar2.text = "2"

        fldEnd = self.create_element('w:fldChar')
        self.create_attribute(fldEnd, 'w:fldCharType', 'end')

        run._r.append(fldStart)
        run._r.append(instrText)
        run._r.append(fldChar1)
        run._r.append(fldChar2)
        run._r.append(fldEnd)

    def find_hyperlinks(self):
        print('Ищу ссылки...')
        self.hyperlinks = []
        run = self.doc.paragraphs[0].runs[0]
        try:
            for bad in run._r.xpath('//w:hyperlink'):
                self.hyperlinks.append((bad, f";{bad.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id']}"))
                print(bad, bad.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'])
                r = self.create_element('w:r')
                t = self.create_element('w:t')
                self.create_attribute(t, 'xml:space', 'preserve')
                t.text = f";{bad.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id']}"
                r.append(t)
                parent = bad.getparent()
                bad.addnext(r)
                parent.remove(bad)
            print('Нашёл ссылки!')
        except Exception as exc:
            print(exc, "find_hyperlinks")

    def revive_hyperlinks(self, path):
        print('Восстанавливаю ссылки...')
        doc = Document(path)
        run = doc.paragraphs[0].add_run()
        for h in self.hyperlinks:
            try:
                for bad in run._r.xpath(f"//w:t[contains(text(),'{h[1]}')]"):
                    text = str(bad.text)
                    textSplitted = text.split(h[1])
                    parent = bad.getparent()
                    if textSplitted[0] == '' and textSplitted.count('') == 1:
                        r = self.create_element('w:r')
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = textSplitted[1]
                        r.append(t)
                        bad.addnext(r)
                        bad.addnext(h[0])
                        parent.remove(bad)
                    elif textSplitted[1] == '' and textSplitted.count('') == 1:
                        r = self.create_element('w:r')
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = textSplitted[0]
                        r.append(t)
                        bad.addnext(h[0])
                        bad.addnext(r)
                        parent.remove(bad)
                    elif textSplitted.count('') > 1:
                        bad.addnext(h[0])
                        parent.remove(bad)
                    elif textSplitted.count('') == 0:
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = textSplitted[0]
                        parent.append(t)
                        r = self.create_element('w:r')
                        r.append(t)
                        r1 = self.create_element('w:r')
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = textSplitted[1]
                        bad.addnext(r1)
                        bad.addnext(h[0])
                        bad.addnext(r)
                        parent.remove(bad)
                print('Восстановил ссылки!')
            except Exception as exc:
                print(exc, "rewrite_highlights")
                continue
        doc.save(path)

    def delete_textInput(self, paragraph):
        run = paragraph.add_run()
        try:
            for bad in run._r.xpath('//w:textInput'):
                bad.getparent().getparent().getparent().getparent().remove(bad.getparent().getparent().getparent())
            for bad in run._r.xpath("//w:instrText[text()=' FORMTEXT ']"):
                bad.getparent().getparent().remove(bad.getparent())
            for bad in run._r.xpath("//w:instrText[text()='FORMTEXT ']"):
                bad.getparent().getparent().remove(bad.getparent())
            for bad in run._r.xpath('//w:smartTag/w:r'):
                anc = bad.getparent()
                anc.addnext(bad)
                anc.getparent().remove(anc)
        except Exception as exc:
            print(exc, "delete_textInput")

    def delete_highlight(self, path):
        doc = Document(path)
        run = doc.paragraphs[0].add_run()
        for bad in run._r.xpath(f"//w:highlight"):
            bad.getparent().remove(bad)
        doc.save(path)

    def find_highlight(self, paragraph, color='yellow'):
        massiv = []
        run = paragraph.add_run()
        for bad in run._r.xpath(f"//w:highlight[@w:val='{color}']"):
            for i in bad.getparent().getparent():
                if i.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                    i.text = self.zamena(i.text)
                    if f"{i.text};hg" not in massiv:
                        massiv.append(f"{i.text};hg" if len(str(i.text).replace(' ', '')) > 1 else f"_{i.text};hg")
                    i.text = f"{i.text};hg" if len(str(i.text).replace(' ', '')) > 1 else f"_{i.text};hg"
        return massiv

    def change_font(self, path):
        doc = Document(path)
        run = doc.paragraphs[0].add_run()
        for bad in run._r.xpath(f"//w:rFonts"):
            try:
                parent = bad.getparent()
                rfont = self.create_element('w:rFonts')
                self.create_attribute(rfont, 'w:ascii', 'Times New Roman')
                self.create_attribute(rfont, 'w:hAnsi', 'Times New Roman')
                self.create_attribute(rfont, 'w:cs', 'Times New Roman')
                self.create_attribute(rfont, 'w:eastAsia', 'Times New Roman')
                parent.remove(bad)
                parent.append(rfont)
            except:
                continue
        doc.save(path)

    def rewrite_highlights(self, hg, path):
        doc = Document(path)
        run = doc.paragraphs[0].add_run()
        for h in hg:
            try:
                for bad in run._r.xpath(f"//w:t[contains(text(),'{h}')]"):
                    text = str(bad.text)
                    textSplitted = text.split(h)
                    parent = bad.getparent()

                    rPr = self.create_element('w:rPr')
                    highlight = self.create_element('w:highlight')
                    self.create_attribute(highlight, 'w:val', 'yellow')
                    rPr.append(highlight)
                    parent.remove(bad)
                    if textSplitted[0] == '' and textSplitted.count('') == 1:
                        parent.append(rPr)
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = h.replace(';hg', '').replace('_', '')
                        parent.append(t)
                        r = self.create_element('w:r')
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = textSplitted[1]
                        r.append(t)
                        parent.addnext(r)
                    elif textSplitted[1] == '' and textSplitted.count('') == 1:
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = textSplitted[0]
                        parent.append(t)
                        r = self.create_element('w:r')
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = h.replace(';hg', '').replace('_', '')
                        r.append(rPr)
                        r.append(t)
                        parent.addnext(r)
                    elif textSplitted.count('') > 1:
                        parent.append(rPr)
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = h.replace(';hg', '').replace('_', '')
                        parent.append(t)
                    elif textSplitted.count('') == 0:
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = textSplitted[0]
                        parent.append(t)
                        r = self.create_element('w:r')
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = h.replace(';hg', '').replace('_', '')
                        r.append(rPr)
                        r.append(t)
                        parent.addnext(r)
                        r1 = self.create_element('w:r')
                        t = self.create_element('w:t')
                        self.create_attribute(t, 'xml:space', 'preserve')
                        t.text = textSplitted[1]
                        r1.append(t)
                        r.addnext(r1)
            except Exception as exc:
                print(exc, "rewrite_highlights : ", h, " : ", bad, " : ", bad.getparent())
                continue
        doc.save(path)

    def zamena(self, text):
        # Убираем двойные пробелы
        if "  " in text:
            text = text.replace('  ', ' ')
        if ' т.е. ' in text:  # проверяем, если ли т.е., если есть заменяем на то есть
            text = text.replace(' т.е. ', ' то есть ')
        if list(self.settings.values())[0]:
            if " N " in text:  # проверяем, если ли N, если есть заменяем на №
                text = text.replace(' N ', ' № ')
        if list(self.settings.values())[2]:
            if '"' in text:  # проверяем, если ли "...", если есть заменяем на «...»
                text = text.replace(' "', ' «')
                text = text.replace('"', '»')
            if '“' in text and '”' in text:  # проверяем, если ли “...”, если есть заменяем на «...»
                text = text.replace('“', '«')
                text = text.replace('”', '»')
        if list(self.settings.values())[4]:
            if ' - ' in text:  # проверяем, если ли -, если есть заменяем на –
                text = text.replace('- ', '– ')
        # Раскрываем аббревиатуры
        if list(self.settings.values())[3]:
            if ' РС (Я) ' in text:
                text = text.replace(' РС (Я) ', ' Республики Саха (Якутия) ')
            if ' РБ ' in text:
                text = text.replace(' РБ ', ' Республики Бурятия ')
            if ' ИО ' in text:
                text = text.replace(' ИО ', ' Иркутской области ')
            if ' ЗК ' in text:
                text = text.replace(' ЗК ', ' Забайкальского края ')
            if 'РС (Я) ' in text:
                text = text.replace('РС (Я) ', 'Республики Саха (Якутия) ')
            if 'РБ ' in text:
                text = text.replace('РБ ', 'Республики Бурятия ')
            if 'ИО ' in text:
                text = text.replace('ИО ', 'Иркутской области ')
            if 'ЗК ' in text:
                text = text.replace('ЗК ', 'Забайкальского края ')
        # Работаем с падежами
        if list(self.settings.values())[5]:
            if 'решение Арбитражный суд' in text:
                text = text.replace('решение Арбитражный суд', 'решение Арбитражного суда')
            if 'Решение Арбитражный суд' in text:
                text = text.replace('Решение Арбитражный суд', 'Решение Арбитражного суда')
            if 'определение Арбитражный суд' in text:
                text = text.replace('определение Арбитражный суд', 'определение Арбитражного суда')
            if 'Определение Арбитражный суд' in text:
                text = text.replace('Определение Арбитражный суд', 'Определение Арбитражного суда')
            if 'наличие в Арбитражный суд' in text:
                text = text.replace('наличие в Арбитражный суд', 'наличие в Арбитражном суде')
            if 'Наличие в Арбитражный суд' in text:
                text = text.replace('Наличие в Арбитражный суд', 'Наличие в Арбитражном суде')
            if 'содействии Арбитражный суд' in text:
                text = text.replace('содействии Арбитражный суд', 'содействии Арбитражного суда')
            if 'Содействии Арбитражный суд' in text:
                text = text.replace('Содействии Арбитражный суд', 'Содействии Арбитражного суда')
            if 'Поручил Арбитражный суд' in text:
                text = text.replace('Поручил Арбитражный суд', 'Поручил Арбитражному суду')
            if 'поручил Арбитражный суд' in text:
                text = text.replace('поручил Арбитражный суд', 'поручил Арбитражному суду')
        # Отдельная опциональная аббревиатура
        if list(self.settings.values())[6]:
            if ' РФ ' in text:
                text = text.replace(' РФ ', ' Российской Федерации ')
        dates, times = get_all_Date_Time(text)
        self.Dates.append(dates)
        self.Times.append(times)
        if list(self.settings.values())[1]:
            if self.Dates[-1]:
                for date in self.Dates[-1]:
                    splittedDate = date.split('.')
                    if len(splittedDate[2]) == 2:
                        text = text.replace(f"{date} г.",
                                            f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                        text = text.replace(date,
                                            f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                    else:
                        text = text.replace(f"{date} г.",
                                            f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                        text = text.replace(date,
                                            f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
        if self.Times[-1]:
            for timen in self.Times[-1]:
                splittedTime = timen.split(':')
                text = text.replace(timen, f"на {splittedTime[0]} часов {splittedTime[1]} минут")
        return text

    # нумерация
    def numbering(self):
        try:
            if self.doc.sections[0].header.paragraphs[0].text == "":
                self.add_page_number(self.doc.sections[0].header.paragraphs[0].add_run())
                self.doc.sections[0].header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # выравниваем по центру
                self.doc.sections[0].different_first_page_header_footer = True  # особый колонтитул для первой страницы - вкл
                sectPr = self.doc.sections[0]._sectPr  # хер его знает, стоило бы узнать
                pgNumType = OxmlElement('w:pgNumType')
                pgNumType.set(ns.qn('w:start'), "1")  # 1 это с какой страницы начинается отсчёт
                sectPr.append(pgNumType)
            return True
        except Exception as exc:
            print(exc, 'numbering')
            return False

    def check_folder_path(self):
        try:
            file_name = os.path.basename(self.path)
            number_doc = file_name[0:file_name.index("_", file_name.index("_") + 1)] if "_" in file_name else file_name
            file_to_save = f"Отформатированный {file_name}"
            if not os.path.exists(self.path_to_save):
                print("Создаю ассистента")
                os.mkdir(self.path_to_save)
            user_file = self.path_to_save + f"\\{number_doc}" if "_" in file_name else self.path_to_save_dif
            if not os.path.exists(user_file):
                print("Создаю пап очка")
                os.mkdir(user_file)
            user_file += "\\" + file_to_save
            return user_file, number_doc, file_name
        except Exception as exc:
            print(exc, 'check_folder_path')

    # Функция для форматирования текст
    def Format(self):
        try:
            # Настройка отступов
            section = self.doc.sections[-1]
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)
            # section.left_margin = Mm(15)
            section.right_margin = Mm(15)
            section.header_distance = Mm(10)
            # отступ от нижнего края страницы до
            # нижнего края нижнего колонтитула
            section.footer_distance = Mm(10)

            print(self.doc.styles)
            for style in self.doc.styles:
                try:
                    style.font.name = 'Times New Roman'
                    style.font.size = Pt(12)
                    # style.font.highlight_color
                except:
                    continue
            # Настройка междустрочного интервала и убираем выделение корректором

            for p in self.doc.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                # p.style.font.highlight_color = WD_COLOR_INDEX.AUTO
            return True
        except Exception as exc:
            print(exc, 'Format')
            return False

    # функция редактирование текста
    def Redact(self):
        print('Нумерую страницы...', end=" ")
        if self.numbering():
            print('Пронумеровал!')
        else:
            print('Не удалось пронумеровать :(')
        print('Форматирую текст...', end=" ")
        if self.Format():
            print('Отформатировал!')
        else:
            print('Не удалось отформатировать :(')

        self.find_hyperlinks()

        if not self.settings["ChangeHighlight"]:
            Highlights = self.find_highlight(self.doc.paragraphs[0])
        else:
            Highlights = []
        for hg in range(len(Highlights)):
            dat, tim = tasklist.get_all_Date_Time(Highlights[hg])
            for d in dat:
                splittedDate = d.split('.')
                if len(splittedDate[2]) == 2:
                    Highlights[hg] = Highlights[hg].replace(f"{d} г.",
                                                            f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                    Highlights[hg] = Highlights[hg].replace(d,
                                                            f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                else:
                    Highlights[hg] = Highlights[hg].replace(f"{d} г.",
                                                            f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                    Highlights[hg] = Highlights[hg].replace(d,
                                                            f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
            for t in tim:
                splittedTime = t.split(':')
                Highlights[hg] = Highlights[hg].replace(t, f"{splittedTime[0]} часов {splittedTime[1]} минут")
            for i in range(10, 99):
                if f"/{i}" in Highlights[hg]:
                    Highlights[hg] = Highlights[hg].replace(f"/{i}", f"/20{i}")
        print(Highlights)
        for p in self.doc.paragraphs:  # проходим все абзацы в документе на поиск ошибок, и заменяем их
            self.delete_textInput(p)
            text = self.zamena(str(p.text))
            if text == "":
                continue
            p.text = text
            # флаг, который отвечает, есть ли ошибка в абзаце
            # если есть, то правим и заменяем текс, если нет, то нет
        if list(self.settings.values())[7]:
            run = self.doc.paragraphs[0].runs[0]
            for i in range(10, 99):
                try:
                    for bad in run._r.xpath(f"//w:t[contains(text(),'/{i}')]"):
                        if bad.text.split(f'/{i}')[1] == '':
                            bad.text = str(bad.text).replace(f'/{i}', f'/20{i}')
                        elif '0' <= bad.text.split(f'/{i}')[1][0] <= '9':
                            continue
                        else:
                            bad.text = str(bad.text).replace(f'/{i}', f'/20{i}')
                except Exception as exc:
                    print(exc)
                    continue
        print(self.Dates)
        print(self.Times)
        path_file, number_file, name_file = self.check_folder_path()
        self.doc.save(path_file)
        if len(Highlights) > 0:
            if self.settings["ChangeHighlight"]:
                self.delete_highlight(path_file)
            else:
                self.rewrite_highlights(Highlights, path_file)
        self.revive_hyperlinks(path_file)
        self.change_font(path_file)
        os.startfile(path_file)
        self.path = path_file
        self.number = number_file
        self.name = name_file
