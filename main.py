from docx import Document
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX

doc = Document('Решение.docx')

month = {"01": "февраля", "02": "января", "03": "марта", "04": "апреля", "05": "мая"
    , "06": "июня", "07": "июля", "08": "августа", "09": "сентября", "10": "октября"
    , "11": "ноября", "12": "декабря"}


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldStart = create_element('w:fldChar')
    create_attribute(fldStart, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'separate')

    fldChar2 = create_element('w:t')
    fldChar2.text = "2"

    fldEnd = create_element('w:fldChar')
    create_attribute(fldEnd, 'w:fldCharType', 'end')

    run._r.append(fldStart)

    run._r.append(instrText)
    run._r.append(fldChar1)
    run._r.append(fldChar2)

    run._r.append(fldEnd)


# нумерация
def numbering():
    add_page_number(doc.sections[0].header.paragraphs[0].add_run())
    doc.sections[0].header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # выравниваем по центру
    doc.sections[0].different_first_page_header_footer = True  # особый колонтитул для первой страницы - вкл
    doc.sections[1].header_distance = Mm(10)  # отступ колонитула от вверхнего края

    sectPr = doc.sections[0]._sectPr  # хер его знает, стоило бы узнать

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "1")  # 1 это с какой страницы начинается отсчёт
    sectPr.append(pgNumType)


# Функция для форматирования текст
def Format():
    # Настройка отступов
    section = doc.sections[-1]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)

    # Настройка междустрочного интервала и убираем выделение корректором
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = Pt(15)
        p.runs[0].font.highlight_color = WD_COLOR_INDEX.AUTO
    # Настройка шрифта и размера текста
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)


# функция редактирование текста
def Redact():
    for p in doc.paragraphs:  # проходим все абзацы в документе на поиск ошибок, и заменяем их
        text = p.text
        # флаг, который отвечает, есть ли ошибка в абзаце
        # если есть, то правим и заменяем текс, если нет, то нет
        flag = False
        if " N " in p.text:  # проверяем, если ли N, если есть заменяем на №
            text = text.replace(' N ', ' № ')
            flag = True
        if ' "' in p.text and '" ' in p.text:  # проверяем, если ли "...", если есть заменяем на «...»
            text = text.replace(' "', ' «')
            text = text.replace('" ', '» ')
            flag = True
        if '“' in p.text and '”' in p.text:  # проверяем, если ли “...”, если есть заменяем на «...»
            text = text.replace('“', ' «')
            text = text.replace('”', '» ')
            flag = True
        if ' - ' in p.text:  # проверяем, если ли -, если есть заменяем на –
            text = text.replace(' - ', ' – ')
            flag = True
        if ' т.е. ' in p.text:  # проверяем, если ли т.е., если есть заменяем на то есть
            text = text.replace(' т.е. ', ' то есть ')
            flag = True
        if flag:  # если есть хотя бы одна ошибка в абзаце, меняем на исправленный вариант
            style = p.style
            p.text = text
            p.style = style
    # редактирование даты, например 12.03.2021 или 12 октября 2021 г.
    # в 12 октября 2021 года
    if "Дело" in doc.paragraphs[5].text:
        text = str(doc.paragraphs[5].text)
        flag = False
        if "." in text[0:10]:
            monthNumb = text[3:5]
            if monthNumb in month.keys():
                text = text.replace(text[0:10], f"{text[0:10]} года")
                text = text.replace(f".{monthNumb}.", f" {month.get(monthNumb)} ")
                flag = True
        elif "г.":
            text = text.replace("г.", "года ")
            flag = True
        if flag:
            style = p.style
            doc.paragraphs[5].text = text
            p.style = style
    doc.save('test.docx')


numbering()
Format()
Redact()
